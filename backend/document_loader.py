"""
Extrae texto de documentos por página/sección y guarda en data/processed_text/ como JSONL.

Uso:
    python document_loader.py
    python document_loader.py --pillar normatividad

Soporta: PDF (PyMuPDF), DOCX (python-docx), HTML (BeautifulSoup), TXT
Salida:  data/processed_text/processed_{pillar}.jsonl
         Cada línea: document_id, page_number, text, status, metadata
"""
import argparse
import json
import logging
import re
from datetime import date
from pathlib import Path

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document

DATA_DIR   = Path(__file__).parent.parent / "data" / "documents"
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "processed_text"
PILARES    = ["normatividad", "costos_operativos", "gestion_activos"]

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _doc_id(path: Path) -> str:
    return re.sub(r"[^a-z0-9]+", "_", path.stem.lower()).strip("_")


def _base_meta(path: Path, pillar: str, file_type: str) -> dict:
    return {
        "pillar":        pillar,
        "title":         path.stem.replace("_", " ").replace("-", " ").title(),
        "filename":      path.name,
        "source":        str(path.relative_to(DATA_DIR.parent.parent)),
        "ingested_date": date.today().isoformat(),
        "file_type":     file_type,
    }


def _ok(doc_id: str, page: int, text: str, meta: dict) -> dict:
    return {"document_id": doc_id, "page_number": page, "text": text, "status": "ok", "metadata": meta}


def _failed(doc_id: str, meta: dict, error: str) -> dict:
    return {"document_id": doc_id, "page_number": None, "text": None,
            "status": "extraction_failed", "error": error, "metadata": meta}


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def load_pdf(path: Path, pillar: str) -> list[dict]:
    doc_id = _doc_id(path)
    meta   = _base_meta(path, pillar, "pdf")
    records = []
    try:
        doc = fitz.open(str(path))
        meta["total_pages"] = len(doc)
        for i, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                records.append(_ok(doc_id, i, text, meta))
        doc.close()
    except Exception as e:
        log.error(f"PDF {path.name}: {e}")
        records.append(_failed(doc_id, meta, str(e)))
    return records


def load_docx(path: Path, pillar: str) -> list[dict]:
    doc_id  = _doc_id(path)
    meta    = _base_meta(path, pillar, "docx")
    records = []
    try:
        doc      = Document(str(path))
        buffer   = []
        page_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            buffer.append(text)
            if len("\n".join(buffer)) >= 800:
                records.append(_ok(doc_id, page_num, "\n".join(buffer), meta))
                buffer    = []
                page_num += 1

        for table in doc.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    for row in table.rows]
            rows = [r for r in rows if r]
            if rows:
                records.append(_ok(doc_id, page_num, "\n".join(rows), meta))
                page_num += 1

        if buffer:
            records.append(_ok(doc_id, page_num, "\n".join(buffer), meta))
    except Exception as e:
        log.error(f"DOCX {path.name}: {e}")
        records.append(_failed(doc_id, meta, str(e)))
    return records


def load_html(path: Path, pillar: str) -> list[dict]:
    doc_id  = _doc_id(path)
    meta    = _base_meta(path, pillar, "html")
    records = []
    try:
        raw  = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if text:
            records.append(_ok(doc_id, 1, text, meta))
    except Exception as e:
        log.error(f"HTML {path.name}: {e}")
        records.append(_failed(doc_id, meta, str(e)))
    return records


def load_txt(path: Path, pillar: str) -> list[dict]:
    doc_id  = _doc_id(path)
    meta    = _base_meta(path, pillar, "txt")
    records = []
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            records.append(_ok(doc_id, 1, text, meta))
    except Exception as e:
        log.error(f"TXT {path.name}: {e}")
        records.append(_failed(doc_id, meta, str(e)))
    return records


LOADERS = {
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".htm":  load_html,
    ".txt":  load_txt,
}


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def process_pillar(pillar: str) -> list[dict]:
    folder = DATA_DIR / pillar
    if not folder.exists():
        log.warning(f"Carpeta no encontrada: {folder}")
        return []

    files = [f for f in folder.iterdir()
             if f.suffix.lower() in LOADERS and not f.name.startswith(".")]
    if not files:
        log.warning(f"Sin documentos en {folder}")
        return []

    all_records = []
    for path in files:
        log.info(f"  → {path.name}")
        records = LOADERS[path.suffix.lower()](path, pillar)
        ok_count = sum(1 for r in records if r["status"] == "ok")
        log.info(f"    {ok_count} páginas/secciones extraídas")
        all_records.extend(records)
    return all_records


def save_jsonl(records: list[dict], pillar: str | None = None) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    suffix   = f"_{pillar}" if pillar else "_all"
    out_path = OUTPUT_DIR / f"processed{suffix}.jsonl"
    with out_path.open("w", encoding="utf-8") as f:
        for rec in records:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    log.info(f"Guardado: {out_path} ({len(records)} registros)")
    return out_path


def run(pillar_filter: str | None = None) -> list[dict]:
    pilares = [pillar_filter] if pillar_filter else PILARES
    all_records = []
    for pillar in pilares:
        log.info(f"\n[{pillar}]")
        all_records.extend(process_pillar(pillar))
    return all_records


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extraer texto de documentos para GuruMaster RAG")
    parser.add_argument("--pillar", choices=PILARES, default=None)
    args = parser.parse_args()

    print("=== GuruMaster — Document Loader (Fase 2) ===")
    records = run(args.pillar)
    save_jsonl(records, args.pillar)

    ok     = sum(1 for r in records if r["status"] == "ok")
    failed = sum(1 for r in records if r["status"] != "ok")
    print(f"\nTotal: {ok} páginas OK | {failed} fallidas")
